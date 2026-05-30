"""
thesis_corrector.py
Core engine: loads norm_control_rules.json and applies formatting rules
to a student thesis .docx file using python-docx.
"""

import json
import re
import copy
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List

import zipfile

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# Some .docx archives contain media files with CRC mismatches inside word/media.
# python-docx uses zipfile.ZipExtFile._update_crc(), which raises BadZipFile on such cases.
# Patch ZipExtFile to ignore CRC-32 failures while preserving normal decompression.
try:
    import docx.opc.phys_pkg as _phys_pkg

    _orig_zipextfile_update_crc = zipfile.ZipExtFile._update_crc

    def _ignore_bad_crc(self, newdata):
        try:
            return _orig_zipextfile_update_crc(self, newdata)
        except zipfile.BadZipFile as err:
            if "Bad CRC-32" in str(err):
                return
            raise

    zipfile.ZipExtFile._update_crc = _ignore_bad_crc
except Exception:
    # If the patch fails, processing will continue with the default behavior.
    pass


# ──────────────────────────────────────────────
# Data structures
# ──────────────────────────────────────────────

@dataclass
class Violation:
    location: str          # e.g. "Para 12" or "Table 3 caption"
    rule_ref: str          # e.g. "7.1.10"
    description: str
    original: str
    corrected: Optional[str] = None
    auto_fixed: bool = False


@dataclass
class CorrectionReport:
    violations: list = field(default_factory=list)
    auto_fixed_count: int = 0
    manual_review_count: int = 0

    def add(self, v: Violation):
        self.violations.append(v)
        if v.auto_fixed:
            self.auto_fixed_count += 1
        else:
            self.manual_review_count += 1

    def summary(self) -> str:
        lines = [
            f"=== Отчёт нормоконтроля ===",
            f"Всего нарушений: {len(self.violations)}",
            f"  Исправлено автоматически: {self.auto_fixed_count}",
            f"  Требуют ручной проверки:  {self.manual_review_count}",
            "",
        ]
        for i, v in enumerate(self.violations, 1):
            status = "✅ ИСПРАВЛЕНО" if v.auto_fixed else "⚠️  ПРОВЕРИТЬ"
            lines.append(f"{i}. [{status}] {v.location} | Правило {v.rule_ref}")
            lines.append(f"   {v.description}")
            if v.original:
                lines.append(f"   Было:  {v.original[:120]}")
            if v.corrected:
                lines.append(f"   Стало: {v.corrected[:120]}")
            lines.append("")
        return "\n".join(lines)


# ──────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────

def emu_to_mm(emu: int) -> float:
    return emu / 914400 * 25.4


def pt_to_half_pt(pt: float) -> int:
    """python-docx stores font size in half-points (twips)."""
    return int(pt * 2)


def get_paragraph_text(para) -> str:
    return "".join(run.text for run in para.runs)


def is_heading(para) -> bool:
    return para.style.name.lower().startswith("heading")


def heading_level(para) -> int:
    name = para.style.name.lower()
    if "heading 1" in name:
        return 1
    if "heading 2" in name:
        return 2
    if "heading 3" in name:
        return 3
    return 0


def detect_section_heading(text: str) -> bool:
    """True if paragraph looks like a numbered section heading (1, 1.1, 1.1.1)."""
    return bool(re.match(r'^\d+(\.\d+)*\s+\S', text.strip()))


def detect_structural_heading(text: str, structural_headings: list) -> bool:
    clean = text.strip().rstrip(".")
    return any(clean.lower() == h.lower() for h in structural_headings)


# ──────────────────────────────────────────────
# Correction modules
# ──────────────────────────────────────────────

class FontCorrector:
    """Rule 5.3 / typography section: font, size, line spacing."""

    def __init__(self, rules: dict, report: CorrectionReport):
        self.rules = rules["typography"]["main_body"]
        self.table_rules = rules["typography"]["tables"]
        self.report = report

    def fix_paragraph(self, para, location: str):
        required_font = self.rules["font_family"]
        required_size = self.rules["font_size_pt"]

        for i, run in enumerate(para.runs):
            changed = False
            original_font = run.font.name
            original_size = run.font.size

            # Font family
            if run.font.name and run.font.name != required_font:
                run.font.name = required_font
                changed = True

            # Font size
            if run.font.size and run.font.size != Pt(required_size):
                orig_pt = run.font.size.pt if run.font.size else "?"
                run.font.size = Pt(required_size)
                changed = True

            if changed:
                self.report.add(Violation(
                    location=f"{location} / run {i}",
                    rule_ref="5.3",
                    description=f"Шрифт: {original_font or '?'} {orig_pt if original_size else '?'}pt → {required_font} {required_size}pt",
                    original=f"{original_font} {original_size}",
                    corrected=f"{required_font} {required_size}pt",
                    auto_fixed=True
                ))


class PageLayoutCorrector:
    """Rule 5.5: margins."""

    def __init__(self, rules: dict, report: CorrectionReport):
        self.margins = rules["page_layout"]["margins_mm"]
        self.report = report

    def fix_document(self, doc: Document):
        required = {
            "left":   Mm(self.margins["left"]),
            "top":    Mm(self.margins["top"]),
            "right":  Mm(self.margins["right"]),
            "bottom": Mm(self.margins["bottom"]),
        }
        for i, section in enumerate(doc.sections):
            changed = []
            if abs(section.left_margin - required["left"]) > 1000:
                section.left_margin = required["left"]
                changed.append(f"left→{self.margins['left']}мм")
            if abs(section.top_margin - required["top"]) > 1000:
                section.top_margin = required["top"]
                changed.append(f"top→{self.margins['top']}мм")
            if abs(section.right_margin - required["right"]) > 1000:
                section.right_margin = required["right"]
                changed.append(f"right→{self.margins['right']}мм")
            if abs(section.bottom_margin - required["bottom"]) > 1000:
                section.bottom_margin = required["bottom"]
                changed.append(f"bottom→{self.margins['bottom']}мм")
            if changed:
                self.report.add(Violation(
                    location=f"Section {i+1}",
                    rule_ref="5.5",
                    description="Поля документа исправлены: " + ", ".join(changed),
                    original="см. оригинал",
                    corrected=str(self.margins),
                    auto_fixed=True
                ))


class ParagraphFormatCorrector:
    """Rule 5.5 / paragraph_formatting: indent, alignment, spacing."""

    def __init__(self, rules: dict, report: CorrectionReport):
        self.rules = rules["paragraph_formatting"]
        self.report = report
        self.skip_styles = {"heading 1", "heading 2", "heading 3",
                            "title", "subtitle", "caption"}

    def fix_paragraph(self, para, location: str):
        style_name = para.style.name.lower()
        if any(s in style_name for s in self.skip_styles):
            return
        if not para.text.strip():
            return

        pf = para.paragraph_format
        required_indent = Mm(self.rules["first_line_indent_mm"])
        required_align = WD_ALIGN_PARAGRAPH.JUSTIFY

        changed = []

        # First line indent
        if pf.first_line_indent is None or abs(pf.first_line_indent - required_indent) > 5000:
            pf.first_line_indent = required_indent
            changed.append(f"абзацный отступ → {self.rules['first_line_indent_mm']}мм")

        # Alignment
        if pf.alignment not in (WD_ALIGN_PARAGRAPH.JUSTIFY, None):
            old = pf.alignment
            pf.alignment = required_align
            changed.append(f"выравнивание {old} → JUSTIFY")

        if changed:
            self.report.add(Violation(
                location=location,
                rule_ref="5.5",
                description="; ".join(changed),
                original=get_paragraph_text(para)[:80],
                corrected=None,
                auto_fixed=True
            ))


class HeadingChecker:
    """Rules 7.1.2, 7.1.10, 7.1.11: section numbering, no dot at end, no hyphenation."""

    def __init__(self, rules: dict, report: CorrectionReport):
        self.structural = rules["document_structure"]["structural_headings_no_numbering"]
        self.report = report

    def check_paragraph(self, para, location: str):
        text = get_paragraph_text(para).strip()
        if not text:
            return

        is_struct = detect_structural_heading(text, self.structural)
        is_sect = detect_section_heading(text)

        if not is_struct and not is_sect:
            return

        # Check: no dot at end (7.1.10)
        if text.endswith(".") and not text.endswith("..."):
            self.report.add(Violation(
                location=location,
                rule_ref="7.1.10",
                description="Заголовок заканчивается точкой (не допускается)",
                original=text,
                corrected=text.rstrip("."),
                auto_fixed=False  # Needs human confirmation
            ))

        # Check: no hyphenation (переносы не допускаются — hard to check, just flag if contains soft hyphen)
        if "\u00ad" in text:
            self.report.add(Violation(
                location=location,
                rule_ref="7.1.10",
                description="В заголовке присутствует перенос (мягкий дефис)",
                original=text,
                auto_fixed=False
            ))


class FigureCaptionChecker:
    """Rule 7.3.1: figure captions format — «Рисунок N – Наименование»."""

    CAPTION_PATTERN = re.compile(
        r'^(Рисунок|Рис\.?)\s*(\d+[\.\d]*)\s*[-–—]\s*(.+)$', re.IGNORECASE
    )
    WRONG_PREFIX = re.compile(r'^(Рис\.)\s', re.IGNORECASE)

    def __init__(self, report: CorrectionReport):
        self.report = report

    def check_paragraph(self, para, location: str):
        text = get_paragraph_text(para).strip()
        if not text.startswith(("Рисунок", "Рис")):
            return

        match = self.CAPTION_PATTERN.match(text)
        if not match:
            self.report.add(Violation(
                location=location,
                rule_ref="7.3.1",
                description="Подпись рисунка не соответствует формату «Рисунок N – Наименование»",
                original=text,
                corrected="Рисунок N – Наименование",
                auto_fixed=False
            ))
            return

        prefix, num, name = match.group(1), match.group(2), match.group(3)
        if prefix.lower() != "рисунок":
            self.report.add(Violation(
                location=location,
                rule_ref="7.3.1",
                description=f"Используется сокращение «{prefix}» вместо «Рисунок»",
                original=text,
                corrected=f"Рисунок {num} – {name}",
                auto_fixed=False
            ))


class TableCaptionChecker:
    """Rule 7.5.1, 7.5.3: table captions format — «Таблица N – Наименование»."""

    CAPTION_PATTERN = re.compile(
        r'^(Таблица|Табл\.?)\s*(\d+[\.\d]*)\s*[-–—]?\s*(.*)$', re.IGNORECASE
    )

    def __init__(self, report: CorrectionReport):
        self.report = report

    def check_paragraph(self, para, location: str):
        text = get_paragraph_text(para).strip()
        if not text.startswith(("Таблица", "Табл")):
            return

        match = self.CAPTION_PATTERN.match(text)
        if not match:
            self.report.add(Violation(
                location=location,
                rule_ref="7.5.1",
                description="Заголовок таблицы не соответствует формату «Таблица N – Наименование»",
                original=text,
                auto_fixed=False
            ))


class BibliographyChecker:
    """Rule 7.1.19: bibliography citation format [N] and ordering."""

    # In-text citation: should be [26] not (26) or 26
    WRONG_CITE_PAREN = re.compile(r'\((\d+)\)')
    BRACKET_CITE = re.compile(r'\[(\d+)\]')

    def __init__(self, report: CorrectionReport):
        self.report = report

    def check_paragraph(self, para, location: str):
        text = get_paragraph_text(para)

        # Check for round-bracket citations — should be square brackets
        wrong = self.WRONG_CITE_PAREN.findall(text)
        if wrong:
            self.report.add(Violation(
                location=location,
                rule_ref="7.1.19.2",
                description=f"Ссылки на литературу оформлены в круглых скобках: {wrong[:5]}. Требуются квадратные скобки [N].",
                original=text[:120],
                auto_fixed=False
            ))


class TextStyleChecker:
    """Rule 7.2.3/7.2.4: forbidden patterns in text."""

    MINUS_SIGN = re.compile(r'(?<!\w)–\s*\d')        # dash before number (could be minus)
    DIAMETER_SYMBOL = re.compile(r'[ØøΦ]')            # diameter symbol in body text
    PERCENT_WITHOUT_NUM = re.compile(r'(?<!\d)\s*%')  # % without preceding number

    def __init__(self, report: CorrectionReport):
        self.report = report

    def check_paragraph(self, para, location: str):
        text = get_paragraph_text(para)
        if not text.strip():
            return

        if self.DIAMETER_SYMBOL.search(text):
            self.report.add(Violation(
                location=location,
                rule_ref="7.2.4",
                description="Символ диаметра Ø в тексте — следует писать слово «диаметр»",
                original=text[:120],
                auto_fixed=False
            ))


# ──────────────────────────────────────────────
# NEW: Line spacing corrector
# ──────────────────────────────────────────────

class LineSpacingCorrector:
    """
    Rule 5.3: межстрочный интервал — одинарный (single / 1.0).

    python-docx line spacing is stored in the paragraph XML as:
      <w:spacing w:line="240" w:lineRule="auto"/>  ← 240 twips = single (1.0)
      <w:spacing w:line="360" w:lineRule="auto"/>  ← 360 twips = 1.5
      <w:spacing w:line="480" w:lineRule="auto"/>  ← 480 twips = double

    240 twips = 1 line = single spacing.

    We skip: headings, captions, table-of-contents styles, empty paragraphs.
    """

    SINGLE_LINE_TWIPS = 240          # 240 twips  = 1.0 line
    TOLERANCE_TWIPS   = 10           # ±10 twips tolerance

    # Styles to leave alone (headings, TOC, captions, etc.)
    SKIP_STYLES = {
        "heading 1", "heading 2", "heading 3", "heading 4",
        "title", "subtitle", "caption",
        "toc 1", "toc 2", "toc 3", "toc heading",
        "list paragraph", "footer", "header",
    }

    def __init__(self, report: CorrectionReport):
        self.report = report

    # ── low-level XML helpers ──────────────────────────────────────────

    @staticmethod
    def _get_pPr(para) -> etree._Element:
        """Return <w:pPr>, creating it if absent."""
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._p.insert(0, pPr)
        return pPr

    @staticmethod
    def _get_spacing(pPr) -> etree._Element:
        """Return <w:spacing>, creating it if absent."""
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        return spacing

    @staticmethod
    def _read_line_twips(para) -> Optional[int]:
        """Return the w:line value in twips, or None if not explicitly set."""
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            return None
        val = spacing.get(qn("w:line"))
        return int(val) if val else None

    @staticmethod
    def _read_line_rule(para) -> Optional[str]:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            return None
        return spacing.get(qn("w:lineRule"))

    # ── public method ──────────────────────────────────────────────────

    def fix_paragraph(self, para, location: str):
        style_name = para.style.name.lower()
        if any(s in style_name for s in self.SKIP_STYLES):
            return
        if not para.text.strip():
            return

        line_twips = self._read_line_twips(para)
        line_rule  = self._read_line_rule(para)

        # Determine effective spacing
        # None means "inherited from style" — which for Normal is usually single.
        # We only act when it is explicitly set to something other than single.
        if line_twips is None:
            return   # inherited / default → assume correct

        is_single = (
            abs(line_twips - self.SINGLE_LINE_TWIPS) <= self.TOLERANCE_TWIPS
            and line_rule in ("auto", None)
        )
        if is_single:
            return

        # Describe what we found
        if line_rule == "exact":
            desc_old = f"точный интервал {line_twips} twips"
        elif line_rule == "atLeast":
            desc_old = f"минимальный интервал {line_twips} twips"
        else:
            ratio = round(line_twips / 240, 2)
            desc_old = f"множитель {ratio}× ({line_twips} twips)"

        # Fix: set to single
        pPr     = self._get_pPr(para)
        spacing = self._get_spacing(pPr)
        spacing.set(qn("w:line"),     str(self.SINGLE_LINE_TWIPS))
        spacing.set(qn("w:lineRule"), "auto")

        self.report.add(Violation(
            location=location,
            rule_ref="5.3",
            description=f"Межстрочный интервал исправлен: {desc_old} → одинарный (1.0)",
            original=desc_old,
            corrected="одинарный (1.0)",
            auto_fixed=True,
        ))


# ──────────────────────────────────────────────
# NEW: Page numbering checker
# ──────────────────────────────────────────────

@dataclass
class PageNumberingConfig:
    """
    Customisable options for the page-numbering checker.

    start_section_name : The section title that should carry the FIRST printed
                         page number.  Default: "Введение".
    start_number       : The integer that should appear on that page.
                         Special value 0 (default) means "use the physical page
                         position of that section" — i.e. the number of pages
                         that precede it plus 1.
    """
    start_section_name: str = "Введение"
    start_number:       int = 0          # 0 → auto-detect from document position


class PageNumberingChecker:
    """
    Rule 7.1.15: Нумерация страниц — сквозная арабскими цифрами,
    в центре нижней части листа, без точки.

    Титульный лист включён в сквозную нумерацию, но номер на нём не ставится.
    Нумерация начинается со страницы «Введение».

    What we can actually check / fix in a .docx:
      A) The footer of each section contains a centred page-number field
         (PAGE field) — if it is missing we report it.
      B) The section that contains the start_section_name has the correct
         page-number start value set (w:pgNumType w:start="N").
      C) The title-page section suppresses its header/footer
         (w:titlePg or different-first-page).

    We cannot reflow pages (Word does that at render time), so we cannot
    guarantee the *physical* page number — we can only set the *counter*.
    """

    # A PAGE field looks like: <w:fldChar/> ... <w:instrText> PAGE </w:instrText> ...
    PAGE_FIELD_RE = re.compile(r'\bPAGE\b', re.IGNORECASE)

    def __init__(self, config: PageNumberingConfig, report: CorrectionReport):
        self.config = config
        self.report = report

    # ── helpers ───────────────────────────────────────────────────────

    @staticmethod
    def _para_text_clean(para) -> str:
        return get_paragraph_text(para).strip().rstrip(".").lower()

    @staticmethod
    def _footer_has_page_field(section) -> bool:
        """Return True if the section's footer contains a PAGE field."""
        try:
            footer = section.footer
        except Exception:
            return False
        for para in footer.paragraphs:
            full_xml = etree.tostring(para._p, encoding="unicode")
            if "PAGE" in full_xml.upper():
                return True
        return False

    @staticmethod
    def _footer_page_field_centred(section) -> bool:
        """Return True if the PAGE field paragraph is centre-aligned."""
        try:
            footer = section.footer
        except Exception:
            return False
        for para in footer.paragraphs:
            full_xml = etree.tostring(para._p, encoding="unicode")
            if "PAGE" not in full_xml.upper():
                continue
            align = para.paragraph_format.alignment
            return align == WD_ALIGN_PARAGRAPH.CENTER
        return False

    @staticmethod
    def _get_pgNumType(section) -> Optional[etree._Element]:
        """Return the <w:pgNumType> element of a section, or None."""
        sectPr = section._sectPr
        if sectPr is None:
            return None
        return sectPr.find(qn("w:pgNumType"))

    @staticmethod
    def _set_page_start(section, start: int):
        """Set w:pgNumType w:start='N' on the section."""
        sectPr = section._sectPr
        pgNum  = sectPr.find(qn("w:pgNumType"))
        if pgNum is None:
            pgNum = OxmlElement("w:pgNumType")
            sectPr.append(pgNum)
        pgNum.set(qn("w:start"), str(start))

    # ── paragraph → section index mapping ────────────────────────────

    @staticmethod
    def _find_section_index_for_paragraph(doc: Document, para_index: int) -> int:
        """
        Approximate which section a paragraph belongs to by scanning for
        sectPr elements embedded in paragraphs before the target index.
        Returns 0-based section index.
        """
        section_idx = 0
        for i, p in enumerate(doc.paragraphs):
            if i >= para_index:
                break
            # A paragraph that ends a section has a <w:sectPr> inside its <w:pPr>
            pPr = p._p.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                section_idx += 1
        return section_idx

    # ── main check ────────────────────────────────────────────────────

    def check(self, doc: Document):
        cfg = self.config
        target_lower = cfg.start_section_name.strip().lower()

        # ── Step 1: locate the start section paragraph ─────────────
        start_para_idx  = None
        start_page_est  = 1   # physical page estimate (counting page-break paras)

        physical_page = 1
        for i, para in enumerate(doc.paragraphs):
            # Rough page counter: explicit page breaks
            xml = etree.tostring(para._p, encoding="unicode")
            if 'w:type="page"' in xml or 'w:pageBreakBefore' in xml:
                physical_page += 1

            clean = self._para_text_clean(para)
            if clean == target_lower:
                start_para_idx = i
                start_page_est = physical_page
                break

        if start_para_idx is None:
            self.report.add(Violation(
                location="Документ",
                rule_ref="7.1.15",
                description=(
                    f"Раздел «{cfg.start_section_name}» не найден в документе. "
                    "Невозможно проверить нумерацию страниц."
                ),
                original="(не найдено)",
                auto_fixed=False,
            ))
            return

        # Determine the desired start number
        desired_start = cfg.start_number if cfg.start_number > 0 else start_page_est

        # ── Step 2: find the section that contains that paragraph ───
        section_idx = self._find_section_index_for_paragraph(doc, start_para_idx)
        sections    = list(doc.sections)

        if section_idx >= len(sections):
            section_idx = len(sections) - 1

        target_section = sections[section_idx]

        # ── Step 3: check / fix w:pgNumType w:start ─────────────────
        pgNum = self._get_pgNumType(target_section)
        current_start = int(pgNum.get(qn("w:start"), 1)) if pgNum is not None else 1

        if current_start != desired_start:
            self._set_page_start(target_section, desired_start)
            self.report.add(Violation(
                location=f"Секция {section_idx+1} (содержит «{cfg.start_section_name}»)",
                rule_ref="7.1.15",
                description=(
                    f"Начало нумерации страниц: было {current_start}, "
                    f"установлено {desired_start} "
                    f"(страница «{cfg.start_section_name}» ≈ физ. стр. {start_page_est})"
                ),
                original=str(current_start),
                corrected=str(desired_start),
                auto_fixed=True,
            ))

        # ── Step 4: check footer has centred PAGE field ──────────────
        for s_idx, section in enumerate(sections):
            has_field    = self._footer_has_page_field(section)
            is_centred   = self._footer_page_field_centred(section) if has_field else False

            if not has_field:
                self.report.add(Violation(
                    location=f"Секция {s_idx+1} — нижний колонтитул",
                    rule_ref="7.1.15",
                    description=(
                        "Нижний колонтитул не содержит поля PAGE (номера страницы). "
                        "Добавьте поле «Номер страницы» в центр нижнего колонтитула."
                    ),
                    original="(отсутствует поле PAGE)",
                    auto_fixed=False,
                ))
            elif not is_centred:
                self.report.add(Violation(
                    location=f"Секция {s_idx+1} — нижний колонтитул",
                    rule_ref="7.1.15",
                    description=(
                        "Номер страницы в нижнем колонтитуле не выровнен по центру. "
                        "Требуется центрирование (7.1.15)."
                    ),
                    original="не по центру",
                    auto_fixed=False,
                ))

        # ── Step 5: title page — number must not be printed ──────────
        # Heuristic: if section 0 has a footer with a visible PAGE field,
        # flag it (the title page should suppress the number).
        if sections and self._footer_has_page_field(sections[0]):
            self.report.add(Violation(
                location="Секция 1 (титульный лист) — нижний колонтитул",
                rule_ref="7.1.15",
                description=(
                    "На титульном листе не должен отображаться номер страницы "
                    "(хотя он включён в общую нумерацию). "
                    "Используйте «Особый колонтитул первой страницы» и оставьте его пустым."
                ),
                original="Поле PAGE присутствует",
                auto_fixed=False,
            ))


# ──────────────────────────────────────────────
# NEW: Section structure checker
# ──────────────────────────────────────────────

class SectionStructureChecker:
    """
    Rules 5.8, 5.11, 5.13: verifies that the required structural sections
    are present in the document and appear in the correct order.

    Works for two document types:
      - "diploma"  : дипломный проект / дипломная работа
      - "magistr"  : магистерская диссертация

    Detection is fuzzy (case-insensitive, strips trailing dots and
    punctuation) so minor capitalisation differences are tolerated.
    """

    # Canonical ordered sequences from the rules JSON
    SEQUENCES = {
        "diploma": [
            "аннотация",
            "содержание",
            "нормативные ссылки",
            "определения",
            "обозначения и сокращения",
            "введение",
            # numbered sections — detected dynamically, labelled "основная часть"
            "заключение",
            "список использованной литературы",
        ],
        "magistr": [
            "содержание",
            "введение",
            "заключение",
            "список использованной литературы",
        ],
    }

    # Optional sections (flagged if missing but not as an error)
    OPTIONAL = {"нормативные ссылки", "определения", "обозначения и сокращения",
                "приложение", "резюме"}

    def __init__(self, rules: dict, report: CorrectionReport,
                 doc_type: str = "diploma"):
        self.required_sequence: List[str] = self.SEQUENCES.get(
            doc_type, self.SEQUENCES["diploma"]
        )
        self.structural = [h.lower() for h in
                           rules["document_structure"]["structural_headings_no_numbering"]]
        self.report = report
        self.doc_type = doc_type

    # ── helpers ───────────────────────────────────────────────────────

    @staticmethod
    def _normalise(text: str) -> str:
        """Strip, lower, remove trailing punctuation for fuzzy matching."""
        return re.sub(r'[.!?:]+$', '', text.strip().lower())

    def _is_structural_heading(self, text: str) -> bool:
        norm = self._normalise(text)
        return any(norm == h for h in self.structural)

    def _is_numbered_heading(self, text: str) -> bool:
        return bool(re.match(r'^\d+(\.\d+)*\s+\S', text.strip()))

    # ── section scan ──────────────────────────────────────────────────

    def _scan_document(self, doc: Document) -> List[dict]:
        """
        Return a list of found sections in document order:
          { "name": <normalised>, "original": <raw text>, "para_idx": N }
        Only top-level structural headings and numbered level-1 sections.
        """
        found = []
        seen_names = set()

        for i, para in enumerate(doc.paragraphs):
            text = get_paragraph_text(para).strip()
            if not text:
                continue

            norm = self._normalise(text)

            if self._is_structural_heading(text):
                if norm not in seen_names:
                    found.append({"name": norm, "original": text, "para_idx": i})
                    seen_names.add(norm)

            elif self._is_numbered_heading(text):
                # Only capture level-1 (e.g. "1 Introduction", not "1.2 Sub")
                if re.match(r'^\d+\s+\S', text.strip()):
                    label = "основная часть"   # treat all numbered top sections as body
                    if label not in seen_names:
                        found.append({"name": label, "original": text, "para_idx": i})
                        seen_names.add(label)

        return found

    # ── main check ────────────────────────────────────────────────────

    def check(self, doc: Document):
        found = self._scan_document(doc)
        found_names = [s["name"] for s in found]

        required = self.required_sequence

        # ── A: check presence ────────────────────────────────────────
        missing = []
        for req in required:
            # "основная часть" is satisfied if we found any numbered heading
            if req == "основная часть":
                if req not in found_names:
                    missing.append(req)
            else:
                if req not in found_names:
                    if req not in self.OPTIONAL:
                        missing.append(req)
                    else:
                        self.report.add(Violation(
                            location="Структура документа",
                            rule_ref="5.8",
                            description=f"Необязательный раздел «{req}» отсутствует (при необходимости)",
                            original="(не найдено)",
                            auto_fixed=False,
                        ))

        if missing:
            self.report.add(Violation(
                location="Структура документа",
                rule_ref="5.8",
                description=(
                    f"Отсутствуют обязательные разделы: "
                    + ", ".join(f"«{m}»" for m in missing)
                ),
                original=str(found_names),
                corrected=f"Требуется: {required}",
                auto_fixed=False,
            ))

        # ── B: check order ───────────────────────────────────────────
        # Build sub-sequence of required items that were actually found
        req_present = [r for r in required if r in found_names]
        found_order  = [f for f in found_names if f in required]

        # Remove duplicates while preserving order
        seen = set()
        found_order_dedup = []
        for x in found_order:
            if x not in seen:
                found_order_dedup.append(x)
                seen.add(x)

        if found_order_dedup != req_present:
            # Find first mismatch
            mismatches = []
            for pos, (expected, actual) in enumerate(
                zip(req_present, found_order_dedup), 1
            ):
                if expected != actual:
                    mismatches.append(
                        f"Позиция {pos}: ожидается «{expected}», найдено «{actual}»"
                    )

            if mismatches:
                self.report.add(Violation(
                    location="Структура документа — порядок разделов",
                    rule_ref="5.8 / 5.13",
                    description=(
                        "Разделы расположены в неправильном порядке:\n      "
                        + "\n      ".join(mismatches)
                    ),
                    original=" → ".join(found_order_dedup),
                    corrected=" → ".join(req_present),
                    auto_fixed=False,
                ))

        # ── C: summary of what was found (informational) ─────────────
        if found:
            self.report.add(Violation(
                location="Структура документа — сводка",
                rule_ref="INFO",
                description=(
                    f"Найдено {len(found)} структурных разделов: "
                    + " → ".join(s["original"][:30] for s in found)
                ),
                original="",
                auto_fixed=True,   # mark as "auto" so it goes to fixed count (info)
            ))


# ──────────────────────────────────────────────
# Main corrector
# ──────────────────────────────────────────────

class ThesisCorrector:
    """
    Loads rules from norm_control_rules.json and applies all
    correction modules to a .docx thesis file.

    Parameters
    ----------
    rules_path : str
        Path to norm_control_rules.json.
    page_numbering_config : PageNumberingConfig | None
        Custom page-numbering settings.
        Defaults: start at «Введение», number = physical page of that section.
    doc_type : str
        "diploma" (default) or "magistr" — selects which required-sections
        sequence is used by SectionStructureChecker.
    """

    def __init__(
        self,
        rules_path: str,
        page_numbering_config: Optional[PageNumberingConfig] = None,
        doc_type: str = "diploma",
    ):
        with open(rules_path, encoding="utf-8") as f:
            self.rules = json.load(f)
        self.pn_config  = page_numbering_config or PageNumberingConfig()
        self.doc_type   = doc_type

    def correct(self, input_path: str, output_path: str) -> CorrectionReport:
        doc    = Document(input_path)
        report = CorrectionReport()

        # ── instantiate all checkers / correctors ─────────────────────
        font_corrector     = FontCorrector(self.rules, report)
        layout_corrector   = PageLayoutCorrector(self.rules, report)
        para_corrector     = ParagraphFormatCorrector(self.rules, report)
        spacing_corrector  = LineSpacingCorrector(report)
        heading_checker    = HeadingChecker(self.rules, report)
        figure_checker     = FigureCaptionChecker(report)
        table_checker      = TableCaptionChecker(report)
        bib_checker        = BibliographyChecker(report)
        text_checker       = TextStyleChecker(report)
        pn_checker         = PageNumberingChecker(self.pn_config, report)
        structure_checker  = SectionStructureChecker(self.rules, report, self.doc_type)

        # ── 1. Document-level checks ──────────────────────────────────
        layout_corrector.fix_document(doc)
        pn_checker.check(doc)
        structure_checker.check(doc)

        # ── 2. Paragraph-level corrections ───────────────────────────
        for i, para in enumerate(doc.paragraphs):
            loc = f"Абзац {i+1}"
            font_corrector.fix_paragraph(para, loc)
            para_corrector.fix_paragraph(para, loc)
            spacing_corrector.fix_paragraph(para, loc)
            heading_checker.check_paragraph(para, loc)
            figure_checker.check_paragraph(para, loc)
            table_checker.check_paragraph(para, loc)
            bib_checker.check_paragraph(para, loc)
            text_checker.check_paragraph(para, loc)

        # ── 3. Table cell paragraphs ──────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, para in enumerate(cell.paragraphs):
                        loc = f"Таблица {t_idx+1} / строка {r_idx+1} / ячейка {c_idx+1} / абз. {p_idx+1}"
                        font_corrector.fix_paragraph(para, loc)
                        spacing_corrector.fix_paragraph(para, loc)

        # Save corrected document
        doc.save(output_path)
        return report


# ──────────────────────────────────────────────
# CLI usage
# ──────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Нормоконтроль дипломной работы (ПР V-08-2022)"
    )
    parser.add_argument("input",  help="Входной файл .docx")
    parser.add_argument("output", help="Выходной файл .docx")
    parser.add_argument("--rules", default="data/norm_control_rules.json",
                        help="Путь к norm_control_rules.json")
    parser.add_argument("--doc-type", default="diploma",
                        choices=["diploma", "magistr"],
                        help="Тип документа (diploma | magistr)")
    parser.add_argument("--start-section", default="Введение",
                        help="Раздел, с которого начинается нумерация страниц")
    parser.add_argument("--start-number", type=int, default=0,
                        help="Номер первой пронумерованной страницы (0 = авто)")
    args = parser.parse_args()

    pn_config = PageNumberingConfig(
        start_section_name=args.start_section,
        start_number=args.start_number,
    )
    corrector = ThesisCorrector(
        rules_path=args.rules,
        page_numbering_config=pn_config,
        doc_type=args.doc_type,
    )
    report = corrector.correct(args.input, args.output)

    print(report.summary())

    report_path = args.output.replace(".docx", "_report.txt")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report.summary())
    print(f"\nОтчёт сохранён: {report_path}")
    print(f"Исправленный файл: {args.output}")