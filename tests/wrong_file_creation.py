from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Wrong margins (default) ──────────────────────
# (left untouched so PageLayoutCorrector fires)

# ── Title page paragraph ─────────────────────────
doc.add_paragraph('ДИПЛОМНЫЙ ПРОЕКТ')

# ── Structural sections in WRONG order (Введение before Содержание) ──
intro = doc.add_paragraph('Введение')   # should come AFTER Содержание
intro.style = doc.styles['Heading 1']

# Body paragraph with 1.5 line spacing (violation)
p = doc.add_paragraph('Актуальность данной работы обусловлена...')
pPr = p._p.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr'); p._p.insert(0, pPr)
spacing = OxmlElement('w:spacing')
spacing.set(qn('w:line'), '360')       # 360 twips = 1.5×
spacing.set(qn('w:lineRule'), 'auto')
pPr.append(spacing)

# Contents section (out of order — after Введение)
toc = doc.add_paragraph('Содержание')
toc.style = doc.styles['Heading 1']

# 1 Основная часть
h1 = doc.add_paragraph('1 Анализ предметной области')
h1.style = doc.styles['Heading 1']

body = doc.add_paragraph('Основной текст первого раздела с двойным интервалом.')
pPr2 = body._p.find(qn('w:pPr'))
if pPr2 is None:
    pPr2 = OxmlElement('w:pPr'); body._p.insert(0, pPr2)
sp2 = OxmlElement('w:spacing')
sp2.set(qn('w:line'), '480')   # 480 = double
sp2.set(qn('w:lineRule'), 'auto')
pPr2.append(sp2)

# Заключение
conc = doc.add_paragraph('Заключение')
conc.style = doc.styles['Heading 1']

# List of references — MISSING «Список использованной литературы»
# (intentional omission to test structure checker)

doc.save('test_full.docx')
print('✅ Тестовый документ создан (test_full.docx)')
print('   Нарушения:')
print('   - Поля: неправильные')
print('   - Межстрочный интервал: 1.5× и 2.0× в двух абзацах')
print('   - Порядок разделов: Введение перед Содержанием')
print('   - Отсутствует: Список использованной литературы, Аннотация')
print('   - Нет колонтитула с нумерацией страниц')
